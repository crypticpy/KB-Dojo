import streamlit as st
import anthropic
import io
import docx
from docx.shared import Inches
from datetime import datetime
import streamlit.components.v1 as components
import json
import os
import base64
import pickle

# Anthropic API configuration
api_key = "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"
client = anthropic.Client(api_key=api_key)


def generate_kb_article(messages):
    try:
        response = client.messages(
            model="claude-v1",
            messages=messages,
            max_tokens=20000,
        )
        return response.content
    except anthropic.APIError as e:
        st.error(f"Error generating KB article: {str(e)}")
        return None


def export_to_word(content, images):
    try:
        doc = docx.Document()
        doc.add_paragraph(content)

        for image in images:
            with io.BytesIO(image["data"]) as stream:
                doc.add_picture(stream, width=Inches(6))

        with io.BytesIO() as output:
            doc.save(output)
            output.seek(0)
            st.download_button(
                label="Export to Word",
                data=output,
                file_name="kb_article.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    except Exception as e:
        st.error(f"Error exporting to Word: {str(e)}")


def load_draft(draft_id):
    try:
        with open(f"drafts/{draft_id}.pkl", "rb") as file:
            return pickle.load(file)
    except FileNotFoundError:
        return None


def save_draft(draft_id, content, images, version):
    try:
        with open(f"drafts/{draft_id}_v{version}.pkl", "wb") as file:
            pickle.dump({"content": content, "images": images}, file)
    except Exception as e:
        st.error(f"Error saving draft: {str(e)}")


def load_version_history(draft_id):
    try:
        versions = []
        for file in os.listdir("drafts"):
            if file.startswith(f"{draft_id}_v") and file.endswith(".pkl"):
                version = file.split("_v")[1][:-4]
                versions.append(version)
        return sorted(versions)
    except Exception as e:
        st.error(f"Error loading version history: {str(e)}")
        return []


def load_draft_version(draft_id, version):
    try:
        with open(f"drafts/{draft_id}_v{version}.pkl", "rb") as file:
            return pickle.load(file)
    except FileNotFoundError:
        return None
    except Exception as e:
        st.error(f"Error loading draft version: {str(e)}")
        return None


def load_templates():
    templates = []
    for file in os.listdir("templates"):
        if file.endswith(".txt"):
            templates.append(file[:-4])
    return templates


def load_template_content(template_name):
    try:
        with open(f"templates/{template_name}.txt", "r") as file:
            return file.read()
    except FileNotFoundError:
        st.error(f"Template '{template_name}' not found.")
        return ""
    except Exception as e:
        st.error(f"Error loading template: {str(e)}")
        return ""


def save_template(template_name, content):
    try:
        with open(f"templates/{template_name}.txt", "w") as file:
            file.write(content)
    except Exception as e:
        st.error(f"Error saving template: {str(e)}")


def analyze_image(image_data):
    try:
        response = client.analyze(
            model="image-analysis",
            image=base64.b64encode(image_data).decode("utf-8"),
        )
        return response.content
    except anthropic.APIError as e:
        st.error(f"Error analyzing image: {str(e)}")
        return None


def main():
    st.set_page_config(page_title="KB DOJO", page_icon=":books:", layout="wide")
    st.title("KB DOJO")

    menu = ["New KB Article", "Load Draft", "Search KB Articles", "Image Analysis"]
    choice = st.sidebar.selectbox("Menu", menu, format_func=lambda
        x: f"{x} {'📝' if x == 'New KB Article' else '📂' if x == 'Load Draft' else '🔍' if x == 'Search KB Articles' else '🖼'}")

    if choice == "New KB Article":
        kb_title = st.text_input("KB Title")
        kb_type = st.selectbox("KB Article Type", ["How-to", "Troubleshooting", "Reference", "FAQ"])
        tags = st.multiselect("Tags", ["Network", "Security", "Hardware", "Software"])

        templates = load_templates()
        selected_template = st.selectbox("Template", ["None"] + templates)

        template_content = ""
        if selected_template != "None":
            template_content = load_template_content(selected_template)

        user_content = st.text_area("User Content", value=template_content, height=200)


        supplemental_content = {}
        num_sections = st.number_input("Number of Supplemental Content Sections", min_value=0, max_value=10, value=0)
        for i in range(num_sections):
            section_name = st.text_input(f"Section {i + 1} Name")
            supplemental_content[section_name] = st.text_area(f"Section {i + 1} Content")

        images = []
        num_images = st.number_input("Number of Images", min_value=0, max_value=5, value=0)
        for i in range(num_images):
            image = st.file_uploader(f"Image {i + 1}", type=["jpg", "jpeg", "png"])
            if image is not None:
                images.append({"name": image.name, "data": image.read()})

        submit_button = st.button("Submit")

        if submit_button:
            if not kb_title.strip():
                st.warning("Please enter a KB title.")
            else:
                if not user_content.strip():
                    st.warning("Please provide some user content.")
                else:
                    with st.spinner("Generating KB article..."):
                        messages = [
                            {"role": "user", "content": f"Title: {kb_title}"},
                            {"role": "user", "content": f"Type: {kb_type}"},
                            {"role": "user", "content": f"Tags: {', '.join(tags)}"},
                            {"role": "user", "content": f"User Content:\n{user_content}"},
                            {"role": "user", "content": "Supplemental Content:"},
                        ]
                    for section_name, section_content in supplemental_content.items():
                        messages.append({"role": "user", "content": f"{section_name}: {section_content}"})

                    for image in images:
                        messages.append({"role": "user", "content": [
                            {
                                "type": "image",
                                "source": {
                                    "type": "base64",
                                    "media_type": "image/jpeg",
                                    "data": image["data"],
                                }
                            },
                            {"type": "text", "text": f"Image: {image['name']}"}
                        ]})

                    result = generate_kb_article(messages)
                    if result:
                        st.markdown(result[0]["text"])

                        export_options = ["Word", "PDF", "HTML", "Markdown"]
                        selected_export = st.selectbox("Export as", export_options)
                        export_button = st.button("Export")
                        if export_button:
                            if selected_export == "Word":
                                export_to_word(result[0]["text"], images)
                            # Add export functionality for other formats

                        save_draft_button = st.button("Save Draft")
                        if save_draft_button:
                            draft_id = f"{kb_title}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                            save_draft(draft_id, result[0]["text"], images, version="1")
                            st.success(f"Draft saved with ID: {draft_id}")

                        feedback = st.text_input("Provide feedback")
                        rating = st.slider("Rate the article", min_value=1, max_value=5, value=3)
                        feedback_button = st.button("Submit Feedback")
                        if feedback_button:
                            # Save feedback and rating
                            pass

    elif choice == "Load Draft":
        draft_id = st.text_input("Draft ID")
        load_button = st.button("Load")
        if load_button:
            version_history = load_version_history(draft_id)
            if version_history:
                selected_version = st.selectbox("Version", version_history, index=len(version_history) - 1)
                draft = load_draft_version(draft_id, selected_version)
                if draft:
                    st.markdown(draft["content"])
                    for image in draft["images"]:
                        st.image(image["data"], caption=image["name"], width=300)

                    export_options = ["Word", "PDF", "HTML", "Markdown"]
                    selected_export = st.selectbox("Export as", export_options)
                    export_button = st.button("Export")
                    if export_button:
                        if selected_export == "Word":
                            export_to_word(draft["content"], draft["images"])
                        # Add export functionality for other formats

                    edit_button = st.button("Edit")
                    if edit_button:
                        # Pre-fill the form fields with the loaded draft data
                        kb_title = draft_id.split("_")[0]
                        user_content = draft["content"]
                        images = draft["images"]

                        # Display the form fields for editing
                        st.text_input("KB Title", value=kb_title)
                        components.html(
                            f'<textarea id="user_content" style="width: 100%; height: 200px;">{user_content}</textarea>',
                            height=200,
                        )

                        num_images = len(images)
                        uploaded_images = []
                        for i in range(num_images):
                            image = st.file_uploader(f"Image {i + 1}", type=["jpg", "jpeg", "png"],
                                                     key=f"edit_image_{i}")
                            if image is not None:
                                uploaded_images.append({"name": image.name, "data": image.read()})

                        save_button = st.button("Save")
                        if save_button:
                            updated_content = components.html(
                                '<script>document.getElementById("user_content").value;</script>'
                            )
                            updated_images = images + uploaded_images
                            new_version = str(int(selected_version) + 1)
                            save_draft(draft_id, updated_content, updated_images, version=new_version)
                            st.success(f"Draft saved as version {new_version}")
                else:
                    st.warning("Draft not found.")
            else:
                st.warning("No versions found for the specified draft ID.")

    elif choice == "Search KB Articles":
        search_query = st.text_input("Search")
        search_button = st.button("Search")
        if search_button:
            # Perform search functionality
            pass

    elif choice == "Image Analysis":
        st.subheader("Image Analysis")
        image = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])
        if image is not None:
            image_data = image.read()
            st.image(image_data, caption=image.name, use_column_width=True)

            analyze_button = st.button("Analyze Image")
            if analyze_button:
                with st.spinner("Analyzing image..."):
                    analysis_result = analyze_image(image_data)
                    if analysis_result:
                        st.markdown(f"**Image Analysis Result:**\n\n{analysis_result[0]['text']}")
                        st.code(analysis_result[0]['text'], language=None)


if __name__ == "__main__":
    main()